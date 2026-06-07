from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Estilos globais ──────────────────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)

# Margens
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color="1F3864"):
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    if level == 1:
        run.font.size = Pt(18)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    return p


def add_para(text, bold=False, italic=False, color=None, indent=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code(text):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    # fundo cinza claro via shading no parágrafo
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F0F0")
    pPr.append(shd)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(2)


def add_table(headers, rows, header_color="1F3864"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalho
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(hdr_cells[i], header_color)

    # Linhas
    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = "FFFFFF" if idx % 2 == 0 else "EBF3FB"
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = cell_text
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(row_cells[j], bg)

    doc.add_paragraph()
    return table


def add_note(text):
    p   = doc.add_paragraph()
    run = p.add_run("Nota: " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x79, 0x55, 0x48)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(6)


def add_divider():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F3864")
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after = Pt(8)


# ════════════════════════════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Guia de Instalação Completo")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor.from_string("1F3864")
title.paragraph_format.space_before = Pt(24)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("FarmTech Solutions — Projeto textinhu")
r2.font.size = Pt(14)
r2.italic = True
r2.font.color.rgb = RGBColor.from_string("2E7D32")

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = sub2.add_run("Passo a passo para configurar o ambiente do zero")
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor.from_string("555555")
sub2.paragraph_format.space_after = Pt(20)

add_divider()

# ════════════════════════════════════════════════════════════════════════════
# AVISO INICIAL
# ════════════════════════════════════════════════════════════════════════════
add_para(
    "Siga os passos na ordem apresentada. Não pule etapas. "
    "Todos os comandos devem ser executados no terminal (PowerShell ou Prompt de Comando) "
    "após cada instalação.",
    italic=True,
    color="B71C1C",
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 1 — PROGRAMAS OBRIGATÓRIOS
# ════════════════════════════════════════════════════════════════════════════
add_heading("1. Programas Obrigatórios", level=1)

# 1.1 Python
add_heading("1.1 Python 3.11", level=2, color="2E7D32")
add_para("Site para download: https://www.python.org/downloads/")
add_bullet('Clique em "Download Python 3.11.x"')
add_bullet('Na tela de instalação, marque obrigatoriamente: "Add Python to PATH"')
add_bullet('Clique em "Install Now" e aguarde a conclusão')
add_para("Verificar instalação — abra o terminal e execute:", bold=True)
add_code("python --version")
add_code("pip --version")
add_para("Ambos devem retornar versões sem mensagem de erro.", italic=True, indent=True)

add_divider()

# 1.2 Git
add_heading("1.2 Git", level=2, color="2E7D32")
add_para("Site para download: https://git-scm.com/download/win")
add_bullet("Execute o instalador e deixe todas as opções no padrão")
add_bullet('Clique "Next" até finalizar')
add_para("Verificar:", bold=True)
add_code("git --version")

add_divider()

# 1.3 VSCode
add_heading("1.3 Visual Studio Code (recomendado)", level=2, color="2E7D32")
add_para("Site para download: https://code.visualstudio.com/")
add_bullet("Execute o instalador")
add_bullet('"Add to PATH" deve estar marcado durante a instalação')
add_para(
    "Extensões obrigatórias dentro do VSCode "
    "(pressione Ctrl+Shift+X, pesquise o nome e clique Install):",
    bold=True,
)
add_table(
    ["Extensão", "ID da Extensão", "Para quê"],
    [
        ["Python",  "ms-python.python",      "Rodar arquivos .py"],
        ["Jupyter", "ms-toolsai.jupyter",    "Abrir notebooks .ipynb"],
        ["R",       "REditorSupport.r",      "Rodar scripts .R"],
    ],
)

add_divider()

# 1.4 R
add_heading("1.4 R 4.x", level=2, color="2E7D32")
add_para("Necessário para FASE 1 e FASE 2 (scripts de análise estatística .R).")
add_para("Site para download: https://cran.r-project.org/bin/windows/base/")
add_bullet("Baixe o instalador R 4.x para Windows")
add_bullet("Execute com todas as opções padrão")
add_para("Verificar (no terminal):", bold=True)
add_code("Rscript --version")

add_divider()

# 1.5 Oracle Instant Client
add_heading("1.5 Oracle Instant Client", level=2, color="2E7D32")
add_para("Necessário APENAS para FASE 2 — módulo de Gestão de Sementes com banco Oracle.")
add_para("Site para download: https://www.oracle.com/database/technologies/instant-client/winx64-64-downloads.html")
add_bullet('Baixe o pacote "Basic Package" (arquivo .zip)')
add_bullet(r'Extraia o conteúdo para: C:\oracle\instantclient')
add_bullet("Adicione esse caminho à variável de ambiente PATH:")
add_bullet('   1. Pesquise "variáveis de ambiente" no menu Iniciar do Windows', )
add_bullet('   2. Em "Variáveis do Sistema" → clique em "Path" → "Editar"')
add_bullet(r'   3. Clique "Novo" → cole: C:\oracle\instantclient')
add_bullet("   4. Clique OK em todas as janelas")
add_bullet("   5. Feche e reabra o terminal")
add_para("Verificar:", bold=True)
add_code('python -c "import oracledb; print(\'Oracle OK\')"')
add_note(
    "Para usar o módulo de sementes você também precisará de acesso a um servidor Oracle "
    "(local ou Oracle Cloud Free Tier). Configure as credenciais no arquivo "
    "FASE2/PythonAlem/database.py após a instalação."
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 2 — CONFIGURAÇÃO DO PROJETO
# ════════════════════════════════════════════════════════════════════════════
add_heading("2. Configuração do Projeto Python", level=1)

add_heading("2.1 Abrir o terminal na pasta do projeto", level=2, color="2E7D32")
add_para(
    "Abra o terminal (PowerShell ou CMD) e navegue até a pasta onde o projeto foi salvo. "
    "Substitua <CAMINHO_DO_PROJETO> pelo caminho real na sua máquina:"
)
add_code(r"cd <CAMINHO_DO_PROJETO>\textinhu")
add_para("Exemplo:", italic=True, indent=True)
add_code(r"cd C:\MeusProjetos\textinhu")

add_divider()

add_heading("2.2 Criar o ambiente virtual", level=2, color="2E7D32")
add_para("Execute o comando abaixo dentro da pasta do projeto:")
add_code("python -m venv .venv")
add_note(
    "Isso cria uma pasta .venv com um Python isolado para o projeto, "
    "evitando conflitos com outros programas instalados na máquina."
)

add_divider()

add_heading("2.3 Ativar o ambiente virtual", level=2, color="2E7D32")
add_para("No PowerShell:", bold=True)
add_code(r".venv\Scripts\Activate.ps1")
add_para("No Prompt de Comando (CMD):", bold=True)
add_code(r".venv\Scripts\activate.bat")
add_para(
    "Após ativar, o terminal exibirá (.venv) no início da linha. "
    "Isso confirma que o ambiente está ativo.",
    italic=True, indent=True
)
add_para("Se aparecer erro de política de execução no PowerShell, execute primeiro:", bold=True)
add_code("Set-ExecutionPolicy -ExecutionPolicy RemoteSigned -Scope CurrentUser")
add_para("Depois ative o ambiente novamente.", indent=True)

add_divider()

add_heading("2.4 Instalar todos os pacotes Python", level=2, color="2E7D32")
add_para("Com o ambiente virtual ativado (você verá (.venv) no terminal):")
add_code("pip install -r requirements.txt")
add_para("Esse comando instala automaticamente todos os pacotes abaixo:", bold=True)
add_table(
    ["Pacote", "Versão Mínima", "Para quê"],
    [
        ["pandas",       "2.0.0",  "Manipulação de dados — todas as fases"],
        ["numpy",        "1.24.0", "Cálculos numéricos"],
        ["matplotlib",   "3.7.0",  "Geração de gráficos"],
        ["seaborn",      "0.12.0", "Gráficos estatísticos"],
        ["scikit-learn", "1.3.0",  "Machine Learning — FASE 3 e 4"],
        ["scipy",        "1.10.0", "Ciência de dados"],
        ["oracledb",     "1.3.0",  "Banco de dados Oracle — FASE 2"],
        ["streamlit",    "1.28.0", "Dashboard web — FASE 4 e 7"],
        ["jupyter",      "1.0.0",  "Notebooks interativos — FASE 5 e 6"],
        ["notebook",     "7.0.0",  "Interface Jupyter no navegador"],
        ["ipykernel",    "6.25.0", "Kernel Python para Jupyter"],
        ["ipython",      "8.14.0", "Terminal Python avançado"],
        ["openpyxl",     "3.1.0",  "Leitura de arquivos .xlsx"],
        ["Pillow",       "10.0.0", "Processamento de imagens"],
        ["tqdm",         "4.65.0", "Barra de progresso"],
        ["PyYAML",       "6.0.0",  "Arquivos de configuração"],
        ["requests",     "2.28.0", "Chamadas HTTP e APIs externas"],
    ],
)

add_divider()

add_heading("2.5 Instalar pacotes R", level=2, color="2E7D32")
add_para("Abra o programa R (ou RStudio se tiver instalado) e execute:")
add_code('install.packages(c("httr", "jsonlite", "ggplot2", "dplyr", "readr"), repos="https://cran.r-project.org")')
add_para("Confirme a instalação quando solicitado.", italic=True, indent=True)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 3 — COMO RODAR O SISTEMA
# ════════════════════════════════════════════════════════════════════════════
add_heading("3. Como Rodar o Sistema", level=1)

add_para(
    "Sempre ative o ambiente virtual antes de rodar qualquer programa "
    "(veja o passo 2.3). "
    "Em seguida, entre na pasta FASE7:",
    bold=True
)
add_code(r"cd <CAMINHO_DO_PROJETO>\textinhu\FASE7")

add_heading("Opção A — Menu interativo no terminal", level=3, color="1565C0")
add_code("python run.py")

add_heading("Opção B — Dashboard visual no navegador", level=3, color="1565C0")
add_code("streamlit run launcher.py")
add_para("O navegador abrirá automaticamente com a interface gráfica.", italic=True, indent=True)

add_heading("Opção C — Fase específica direto pela linha de comando", level=3, color="1565C0")
add_table(
    ["Comando", "O que faz"],
    [
        ["python run.py 1",         "Gestão de Áreas e Insumos (FASE 1)"],
        ["python run.py 2a",        "Gestão de Sementes com Oracle (FASE 2)"],
        ["python run.py 2b",        "Consulta Meteorológica via API (FASE 1 R)"],
        ["python run.py 2c",        "Análise R — Áreas por Cultura (FASE 1 R)"],
        ["python run.py 3",         "Classificação de Culturas com ML (FASE 3)"],
        ["python run.py 4t",        "Treinar Modelos IoT (FASE 4)"],
        ["python run.py 4d",        "Dashboard ML — Irrigação (FASE 4)"],
        ["python run.py 5",         "Previsão de Rendimento — Jupyter (FASE 5)"],
        ["python run.py 6",         "Visão Computacional YOLO — Jupyter (FASE 6)"],
        ["python run.py --list",    "Listar todas as fases disponíveis"],
        ["python run.py dashboard", "Abrir o dashboard visual (launcher)"],
    ],
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 4 — GOOGLE COLAB (FASE 5 e 6)
# ════════════════════════════════════════════════════════════════════════════
add_heading("4. FASE 5 e FASE 6 — Alternativa via Google Colab", level=1)
add_para(
    "As FASE 5 (Previsão de Rendimento) e FASE 6 (Visão Computacional com YOLO) "
    "utilizam notebooks Jupyter (.ipynb). Você tem duas opções:"
)
add_bullet("Rodar localmente: use o comando python run.py 5 ou python run.py 6 (Jupyter abre no navegador)")
add_bullet(
    "Rodar no Google Colab (sem instalação): acesse https://colab.research.google.com, "
    'clique em "Fazer upload de notebook" e selecione o arquivo .ipynb da fase correspondente. '
    "Nesse caso, não é necessária nenhuma instalação local para essas fases."
)

# ════════════════════════════════════════════════════════════════════════════
# SEÇÃO 5 — RESUMO
# ════════════════════════════════════════════════════════════════════════════
add_heading("5. Resumo — Lista de Downloads", level=1)
add_table(
    ["#", "Programa", "Site", "Obrigatório para"],
    [
        ["1", "Python 3.11",           "python.org/downloads",                   "Tudo"],
        ["2", "Git",                   "git-scm.com/download/win",               "Clonar o repositório"],
        ["3", "Visual Studio Code",    "code.visualstudio.com",                  "Editor (recomendado)"],
        ["4", "R 4.x",                 "cran.r-project.org/bin/windows/base",    "FASE 1 e 2 (scripts .R)"],
        ["5", "Oracle Instant Client", "oracle.com → Instant Client Downloads",  "Apenas FASE 2"],
        ["6", "Ext. VSCode: Python",   "Dentro do VSCode (Ctrl+Shift+X)",        "Editar/rodar .py"],
        ["7", "Ext. VSCode: Jupyter",  "Dentro do VSCode (Ctrl+Shift+X)",        "Abrir .ipynb"],
        ["8", "Ext. VSCode: R",        "Dentro do VSCode (Ctrl+Shift+X)",        "Editar/rodar .R"],
    ],
)

# ════════════════════════════════════════════════════════════════════════════
# RODAPÉ
# ════════════════════════════════════════════════════════════════════════════
add_divider()
rodape = doc.add_paragraph()
rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = rodape.add_run("FarmTech Solutions — Projeto IA_Underground — FIAP ON")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x90, 0x90, 0x90)

# ════════════════════════════════════════════════════════════════════════════
# SALVAR
# ════════════════════════════════════════════════════════════════════════════
output = r"C:\Users\josel\Downloads\textinhu\textinhu\Guia_Instalacao_FarmTech.docx"
doc.save(output)
print(f"Documento criado: {output}")
